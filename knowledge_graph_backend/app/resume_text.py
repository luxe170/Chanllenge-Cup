"""Extract plain text from uploaded resume files (PDF / DOCX / TXT).

PDF extraction order:
    1. PyMuPDF (fitz) — best layout & Chinese fidelity when available;
    2. pypdf — pure-python fallback if PyMuPDF import fails;
    3. OCR — if the text-layer extractors yield too little text (e.g. scanned
       resumes), each page is rendered to PNG and forwarded to the configured
       :class:`~app.ocr.OcrProvider`. If no provider is configured the file is
       rejected with a clear error.
"""
from __future__ import annotations

import io
import logging
import unicodedata
from typing import Final

from .ocr import OcrError, OcrProvider, get_ocr_provider


logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS: Final[tuple[str, ...]] = (".pdf", ".doc", ".docx", ".txt", ".md")

# When text-layer extraction returns fewer than this many non-whitespace
# characters we treat the PDF as scanned and hand it to OCR.
MIN_TEXT_CHARS: Final[int] = 40


class ResumeTextError(ValueError):
    """Raised when a resume file cannot be parsed into text."""


def _normalize(text: str) -> str:
    text = unicodedata.normalize("NFKC", text or "")
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def _text_length(text: str) -> int:
    return sum(1 for ch in (text or "") if not ch.isspace())


def _read_pdf_pymupdf(data: bytes) -> str:
    try:
        import fitz  # type: ignore  # PyMuPDF
    except ImportError:
        return ""
    try:
        document = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:  # PyMuPDF raises RuntimeError / ValueError on bad PDFs
        raise ResumeTextError(f"invalid pdf file: {exc}") from exc
    try:
        parts: list[str] = []
        for page in document:
            try:
                parts.append(page.get_text("text") or "")
            except Exception:  # noqa: BLE001 - a single bad page shouldn't kill the file
                parts.append("")
        return "\n".join(parts)
    finally:
        document.close()


def _read_pdf_pypdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ResumeTextError("no PDF backend is available") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ResumeTextError(f"invalid pdf file: {exc}") from exc
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            parts.append("")
    return "\n".join(parts)


def _pdf_page_images(data: bytes, *, dpi: int = 220) -> list[bytes]:
    """Rasterize each page to a PNG for OCR. Returns an empty list if PyMuPDF is missing."""

    try:
        import fitz  # type: ignore
    except ImportError:  # pragma: no cover - OCR path requires PyMuPDF
        return []
    try:
        document = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise ResumeTextError(f"invalid pdf file: {exc}") from exc
    try:
        images: list[bytes] = []
        matrix = fitz.Matrix(dpi / 72, dpi / 72)
        for page in document:
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            images.append(pixmap.tobytes("png"))
        return images
    finally:
        document.close()


def _ocr_pdf(data: bytes, provider: OcrProvider) -> str:
    if not provider.available:
        raise ResumeTextError(
            "PDF has no extractable text and no OCR provider is configured. "
            "Set OCR_PROVIDER and OCR_ENDPOINT, or upload a text-based PDF/DOCX."
        )
    images = _pdf_page_images(data)
    if not images:
        raise ResumeTextError(
            "PDF appears to be scanned but page rendering is unavailable "
            "(PyMuPDF is required for OCR)."
        )
    parts: list[str] = []
    for index, image in enumerate(images, start=1):
        try:
            parts.append(provider.recognize(image, "image/png"))
        except OcrError as exc:
            logger.warning("OCR failed on page %s: %s", index, exc)
    combined = "\n".join(part for part in parts if part)
    if not combined.strip():
        raise ResumeTextError("OCR did not return any text for this resume")
    return combined


def _read_pdf(data: bytes) -> str:
    text = _read_pdf_pymupdf(data)
    if _text_length(text) < MIN_TEXT_CHARS:
        # PyMuPDF may return partial text on some layouts; give pypdf a chance.
        fallback = _read_pdf_pypdf(data)
        if _text_length(fallback) > _text_length(text):
            text = fallback
    if _text_length(text) < MIN_TEXT_CHARS:
        # Scanned PDF (or a text layer we cannot read). Route to OCR.
        text = _ocr_pdf(data, get_ocr_provider())
    return text


def _read_docx(data: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise ResumeTextError("python-docx is required to parse DOCX resumes") from exc
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ResumeTextError(f"invalid docx file: {exc}") from exc
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _read_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ResumeTextError("resume text encoding not recognized")


def extract_resume_text(filename: str, data: bytes) -> str:
    """Return normalized text extracted from ``data``, dispatched by extension."""

    if not data:
        raise ResumeTextError("empty file")
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        text = _read_pdf(data)
    elif lower.endswith(".docx"):
        text = _read_docx(data)
    elif lower.endswith(".doc"):
        raise ResumeTextError("legacy .doc files are not supported; please upload .docx or .pdf")
    elif lower.endswith((".txt", ".md")):
        text = _read_plain(data)
    else:
        raise ResumeTextError(f"unsupported resume format: {filename}")
    normalized = _normalize(text)
    if not normalized:
        raise ResumeTextError("could not extract any text from the resume")
    return normalized
