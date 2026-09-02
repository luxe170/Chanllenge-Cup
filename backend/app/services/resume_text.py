from __future__ import annotations

import io
import unicodedata
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Final
from xml.etree import ElementTree

from backend.app.services.ocr import OcrError, OcrProvider, get_ocr_provider


SUPPORTED_EXTENSIONS: Final[tuple[str, ...]] = (".pdf", ".doc", ".docx", ".txt", ".md", ".png", ".jpg", ".jpeg", ".webp")
MIN_TEXT_CHARS: Final[int] = 40
MAX_VISION_PAGES: Final[int] = 8
VISION_RENDER_DPI: Final[int] = 160


class ResumeTextError(ValueError):
    """Raised when a resume file cannot be parsed into usable text."""


@dataclass(slots=True)
class ResumeContent:
    """Text or page images ready for the resume analyzer."""

    text: str = ""
    images: list[bytes] = field(default_factory=list)
    mode: str = "text"
    mime_type: str = "image/png"


def _normalize(text: str) -> str:
    text = unicodedata.normalize("NFKC", text or "")
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def _text_length(text: str) -> int:
    return sum(1 for char in text if not char.isspace())


def _read_pdf_pymupdf(data: bytes) -> str:
    try:
        import fitz  # type: ignore
    except ImportError:
        return ""
    try:
        document = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise ResumeTextError(f"PDF 文件无效或已损坏: {exc}") from exc
    try:
        parts = []
        for page in document:
            try:
                parts.append(page.get_text("text") or "")
            except Exception:
                parts.append("")
        return "\n".join(parts)
    finally:
        document.close()


def _read_pdf_pypdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ResumeTextError("缺少 PDF 解析依赖 pypdf") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ResumeTextError(f"PDF 文件无效或已损坏: {exc}") from exc
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts)


def _pdf_page_images(data: bytes, *, dpi: int = 220) -> list[bytes]:
    try:
        import fitz  # type: ignore
    except ImportError:
        return []
    try:
        document = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise ResumeTextError(f"PDF 文件无效或已损坏: {exc}") from exc
    try:
        matrix = fitz.Matrix(dpi / 72, dpi / 72)
        return [page.get_pixmap(matrix=matrix, alpha=False).tobytes("png") for page in document]
    finally:
        document.close()


def _ocr_pdf(data: bytes, provider: OcrProvider) -> str:
    if not provider.available:
        raise ResumeTextError("扫描版 PDF 未检测到可抽取文本，且 OCR 未配置。请设置 OCR_PROVIDER=http 和 OCR_ENDPOINT，或上传文字版 PDF/DOCX。")
    images = _pdf_page_images(data)
    if not images:
        raise ResumeTextError("扫描版 PDF 需要 PyMuPDF 渲染页面后才能 OCR。")
    parts = []
    for image in images:
        try:
            text = provider.recognize(image, "image/png")
        except OcrError:
            text = ""
        if text:
            parts.append(text)
    combined = "\n".join(parts)
    if not combined.strip():
        raise ResumeTextError("OCR 未返回可用文本。")
    return combined


def _read_pdf(data: bytes) -> str:
    text = _read_pdf_text(data)
    if _text_length(text) < MIN_TEXT_CHARS:
        text = _ocr_pdf(data, get_ocr_provider())
    return text


def _read_pdf_text(data: bytes) -> str:
    text = _read_pdf_pymupdf(data)
    if _text_length(text) < MIN_TEXT_CHARS:
        fallback = _read_pdf_pypdf(data)
        if _text_length(fallback) > _text_length(text):
            text = fallback
    return text


def _read_docx_with_python_docx(data: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return ""
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ResumeTextError(f"Word 文件无法解析或已经损坏: {exc}") from exc
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _read_docx_xml(data: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            xml = archive.read("word/document.xml")
    except Exception as exc:
        raise ResumeTextError("Word 文件无法解析或已经损坏") from exc
    root = ElementTree.fromstring(xml)
    return "\n".join("".join(node.itertext()) for node in root.iter() if node.tag.endswith("}p"))


def _read_docx(data: bytes) -> str:
    text = _read_docx_with_python_docx(data)
    return text if text.strip() else _read_docx_xml(data)


def _read_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ResumeTextError("简历文本编码无法识别。")


def extract_resume_text(filename: str, content: bytes) -> str:
    if not content:
        raise ResumeTextError("上传的简历为空")
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf(content)
    elif suffix == ".docx":
        text = _read_docx(content)
    elif suffix == ".doc":
        raise ResumeTextError("暂不支持旧版 .doc，请另存为 .docx 或 PDF")
    elif suffix in {".txt", ".md"}:
        text = _read_plain(content)
    else:
        raise ResumeTextError("仅支持 PDF、DOCX、TXT 简历")
    normalized = _normalize(text)
    if not normalized:
        raise ResumeTextError("没有从简历中提取到可用文本")
    return normalized


def extract_resume_content(filename: str, content: bytes, *, allow_vision: bool = False) -> ResumeContent:
    """Prepare a resume for analysis without changing the public result schema.

    Text-bearing files stay on the text path. When a PDF has insufficient embedded
    text and multimodal analysis is enabled, its pages are rendered as images and
    returned directly instead of being OCR'd first.
    """
    if not content:
        raise ResumeTextError("上传的简历为空")
    suffix = Path(filename or "").suffix.lower()
    if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
        if not allow_vision:
            raise ResumeTextError("图片简历需要启用支持视觉输入的简历分析模型。")
        mime_type = "image/jpeg" if suffix in {".jpg", ".jpeg"} else f"image/{suffix.lstrip('.')}"
        return ResumeContent(images=[content], mode="vision", mime_type=mime_type)
    if suffix != ".pdf" or not allow_vision:
        return ResumeContent(text=extract_resume_text(filename, content))

    text = _normalize(_read_pdf_text(content))
    if _text_length(text) >= MIN_TEXT_CHARS:
        return ResumeContent(text=text)

    images = _pdf_page_images(content, dpi=VISION_RENDER_DPI)
    if not images:
        raise ResumeTextError("扫描版 PDF 无法渲染为图片，请安装 PyMuPDF 或检查文件是否损坏。")
    if len(images) > MAX_VISION_PAGES:
        raise ResumeTextError(f"扫描版 PDF 超过 {MAX_VISION_PAGES} 页，请精简后重新上传。")
    return ResumeContent(images=images, mode="vision")
